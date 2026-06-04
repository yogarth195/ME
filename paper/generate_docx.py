"""Generate ExplainHire Word (.docx) report from scratch using python-docx."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

# ── Style helpers ─────────────────────────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(text, bold=False, italic=False, center=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_placeholder(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ Figure Placeholder: {label} ]')
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    border = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), 'AAAAAA')
        border.append(b)
    p._p.get_or_add_pPr().append(border)
    doc.add_paragraph()

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_para('DEPT. OF COMPUTER SCIENCE & ENGINEERING', bold=True, center=True, size=14)
add_para('DELHI TECHNOLOGICAL UNIVERSITY', bold=True, center=True, size=14)
add_para('(Formerly Delhi College of Engineering)', center=True, size=11)
add_para('Bawana Road, Delhi-110042', center=True, size=11)
doc.add_paragraph()
add_para('─' * 60, center=True, size=11)
doc.add_paragraph()
add_para('ExplainHire: Neurosymbolic AI for', bold=True, center=True, size=18)
add_para('Explainable Resume–Job Matching', bold=True, center=True, size=18)
doc.add_paragraph()
add_para('─' * 60, center=True, size=11)
doc.add_paragraph()
add_para('A PROJECT REPORT', bold=True, center=True, size=13)
add_para('Submitted in Partial Fulfillment of the Requirements\nfor the Award of the Degree of', italic=True, center=True)
add_para('BACHELOR OF TECHNOLOGY\nIN\nCOMPUTER SCIENCE & ENGINEERING', bold=True, center=True, size=14)
doc.add_paragraph()
add_para('Submitted by:', bold=True, center=True)
add_para('Yoosha Raza  (2K22/CO/521)', center=True)
add_para('Yanshu  (2K22/CO/507)', center=True)
add_para('Yogarth  (2K22/CO/519)', center=True)
doc.add_paragraph()
add_para('Under the supervision of', bold=True, center=True)
add_para('Gull Kaur', bold=True, center=True)
add_para('Assistant Professor\nDepartment of Computer Science & Engineering\nDelhi Technological University', center=True)
doc.add_paragraph()
add_para('MAY 2026', bold=True, center=True, size=14)
doc.add_page_break()

# ── Declaration ───────────────────────────────────────────────────────────────
add_para('CANDIDATE\'S DECLARATION', bold=True, center=True, size=16)
doc.add_paragraph()
doc.add_paragraph(
    "We, Yoosha Raza (2K22/CO/521), Yanshu (2K22/CO/507), and Yogarth (2K22/CO/519), "
    "students of B.Tech (Computer Science & Engineering), Batch 2022–2026, hereby declare "
    "that the Project Dissertation titled \"ExplainHire: Neurosymbolic AI for Explainable "
    "Resume–Job Matching\", submitted to CSE, DTU, is an original work carried out by us. "
    "This work has not previously formed the basis for the award of any Degree or Diploma. "
    "All sources have been duly cited.\n\n"
    "Place: New Delhi       Date: May 2026"
)
doc.add_paragraph()
t = doc.add_table(rows=2, cols=3)
t.cell(0, 0).text = 'Yoosha Raza'
t.cell(0, 1).text = 'Yanshu'
t.cell(0, 2).text = 'Yogarth'
t.cell(1, 0).text = '(2K22/CO/521)'
t.cell(1, 1).text = '(2K22/CO/507)'
t.cell(1, 2).text = '(2K22/CO/519)'
doc.add_page_break()

# ── Certificate ───────────────────────────────────────────────────────────────
add_para('CERTIFICATE', bold=True, center=True, size=16)
doc.add_paragraph()
doc.add_paragraph(
    "This is to certify that the project report titled \"ExplainHire: Neurosymbolic AI for "
    "Explainable Resume–Job Matching\", submitted by Yoosha Raza (2K22/CO/521), "
    "Yanshu (2K22/CO/507), and Yogarth (2K22/CO/519) in partial fulfillment of the "
    "requirements for B.Tech in CSE from DTU, is a bonafide record of original project work "
    "carried out under my guidance during 2025–26. I recommend this report for acceptance.\n\n"
    "Place: New Delhi       Date: May 2026"
)
doc.add_paragraph()
doc.add_paragraph("Gull Kaur\nAssistant Professor, CSE, DTU, Delhi-110042")
doc.add_page_break()

# ── Abstract ──────────────────────────────────────────────────────────────────
add_para('ABSTRACT', bold=True, center=True, size=16)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Keywords — ').bold = True
p.add_run('Neurosymbolic AI, Resume–Job Matching, Skill Ontology, Sentence-BERT, XGBoost, SHAP Explainability, NLP')
doc.add_paragraph(
    "Existing Applicant Tracking Systems (ATS) rely on keyword co-occurrence to screen "
    "resumes, causing semantic blindness (equivalent skills described differently score zero) "
    "and decision opacity (no structured justification for outcomes). This report presents "
    "ExplainHire, a neurosymbolic AI system that resolves both failures."
)
doc.add_paragraph(
    "ExplainHire integrates three complementary signals in a seven-layer pipeline: "
    "(1) a 434-node symbolic skill ontology graph from O*NET with six-level weighted traversal; "
    "(2) a domain-adapted SBERT model fine-tuned on 887 real resume–JD pairs; and "
    "(3) structural features (YOE gap, education level, section coverage). "
    "These yield 14 numeric features fused by XGBoost with SHAP explainability and "
    "sentence-level evidence mapping."
)
doc.add_paragraph(
    "ExplainHire achieves 5-fold CV F1 = 0.9045 and full-dataset F1 = 0.951, "
    "+11.86% over the best single-signal baseline. On the public ResumeAtlas benchmark "
    "it achieves F1 = 0.720 vs. 0.610 reported baseline (+11.06%). It is the first "
    "resume–job matching system to provide structured, auditable, sentence-level "
    "explanations with ranked skill gap analysis."
)
doc.add_page_break()

# ── Acknowledgement ───────────────────────────────────────────────────────────
add_para('ACKNOWLEDGEMENT', bold=True, center=True, size=16)
doc.add_paragraph()
doc.add_paragraph(
    "We are deeply grateful to our supervisor, Gull Kaur, Assistant Professor, CSE, DTU, "
    "for her invaluable guidance throughout this project. We thank the Head of Department "
    "and all CSE faculty at DTU. We owe a special debt to the 144 fellow B.Tech students "
    "who generously shared their resumes — our results rest on their participation. "
    "We also thank the creators of O*NET, sentence-transformers, ResumeAtlas, and the "
    "Kaggle resume dataset. Finally, we thank our families and parents whose unwavering "
    "support through four years made everything possible."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Chapter 1: INTRODUCTION', level=1)

add_heading('1.1 Overview and Motivation', level=2)
doc.add_paragraph(
    "The modern job market operates at a scale that exceeds human processing capacity. "
    "A single corporate job posting routinely receives hundreds of applications within 48 hours, "
    "making automated resume screening operationally necessary. The dominant technology is the "
    "Applicant Tracking System (ATS) — platforms such as Workday, Greenhouse, and their Indian "
    "equivalents in Naukri and LinkedIn. Industry surveys show 75–88% of applicants are "
    "eliminated by ATS before any human reads their application."
)
doc.add_paragraph(
    "This approach suffers from two fundamental failures: semantic blindness (a system scoring "
    "'built distributed backend services with Go' as zero against 'Golang developer with "
    "microservices experience' has only counted words, not understood meaning) and decision "
    "opacity (neither candidates nor recruiters can obtain any structured justification). "
    "Emerging regulations such as the EU AI Act classify employment AI as 'high-risk' and "
    "mandate explainability and audit trails."
)
doc.add_paragraph(
    "ExplainHire is a neurosymbolic AI system combining a 434-node skill ontology graph "
    "(symbolic AI), a fine-tuned SBERT sentence embedding model (neural AI), and XGBoost "
    "(learned fusion) — delivering match/no-match predictions with full sentence-level, "
    "SHAP-attributed explanations for every decision."
)

add_heading('1.2 Problem Statement', level=2)
doc.add_paragraph(
    "Given a resume R (PDF/DOCX) and job description J (plain text), produce: "
    "(1) a binary label ŷ ∈ {0,1} with confidence p ∈ [0,1]; "
    "(2) an explanation containing per-skill match evidence (matched skill, ontological distance, "
    "resume sentence), ranked missing skills by counterfactual impact, and SHAP feature attribution "
    "across 14 input features."
)

add_heading('1.3 Research Contributions', level=2)
add_numbered('Neurosymbolic matching pipeline: first system combining symbolic ontology traversal, fine-tuned neural embeddings, and structural features fused by a learned classifier.')
add_numbered('Domain-specific skill ontology: 434-node directed graph from O*NET, extended for Indian software engineering hiring, with an 800-entry alias normalization table.')
add_numbered('Domain-adapted SBERT: all-MiniLM-L6-v2 fine-tuned on 887 real resume–JD pairs.')
add_numbered('Structured explainability: sentence-level evidence mapping and counterfactual skill gap ranking.')
add_numbered('State-of-the-art results: CV F1 = 0.9045, full-data F1 = 0.951, ResumeAtlas F1 = 0.720.')

add_heading('1.4 Scope and Limitations', level=2)
doc.add_paragraph(
    "ExplainHire is validated for technical and software engineering roles in English. "
    "Extension to non-technical domains, multi-lingual resumes, and multi-candidate ranking "
    "are future directions."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Chapter 2: BACKGROUND AND RELATED WORK', level=1)

add_heading('2.1 The Recruitment Technology Landscape', level=2)
doc.add_paragraph(
    "Despite filtering 75–88% of applicants, the core ATS algorithm is primitive: parse resume "
    "into fields, extract JD keywords, score by keyword overlap, rank and cutoff. This is "
    "lexically sensitive but semantically blind — 'developed machine learning models for fraud "
    "detection' scores zero against 'AI engineer with anomaly detection experience'."
)

add_heading('2.2 Prior Work', level=2)
doc.add_paragraph(
    "PJFNN (Zhu et al., 2018) used hierarchical CNN on Baidu click logs, F1=0.800. "
    "Requires large implicit feedback data, no explanations, ignores skill structure."
)
doc.add_paragraph(
    "Li et al. (2020) applied BERT + multi-head attention, accuracy 0.792. Semantically richer "
    "but still a black box."
)
doc.add_paragraph(
    "Bian et al. (2020) combined GNN + BERT multi-view co-teaching, ~F1=0.78. "
    "Graph is co-occurrence-based (not a domain ontology); no explainability."
)
doc.add_paragraph(
    "conSultantBERT (Lavi et al., 2021) fine-tuned SBERT on Randstad data, F1=0.749. "
    "Closest precedent to our SBERT component; uses proprietary data; no symbolic or structural signals."
)
doc.add_paragraph(
    "Heakl et al. (2024) introduced ResumeAtlas (13,389 resumes, 43 categories) with "
    "TF-IDF + XGBoost baseline F1=0.610 — the only public large-scale benchmark."
)

add_heading('2.3 Comparison Table', level=2)
t = doc.add_table(rows=1, cols=6)
t.style = 'Table Grid'
hdr = t.rows[0].cells
for i, h in enumerate(['Feature', 'PJFNN', 'Li\'20', 'Bian\'20', 'cBERT', 'ExplainHire']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ['Structured ontology', '✗', '✗', 'partial', '✗', '✓'],
    ['Fine-tuned SBERT', '✗', '✗', '✗', '✓', '✓'],
    ['Structural features', '✗', '✗', '✗', '✗', '✓'],
    ['Decision explanation', '✗', '✗', '✗', '✗', '✓'],
    ['Skill gap analysis', '✗', '✗', '✗', '✗', '✓'],
    ['Public validation', '✗', '✗', '✗', '✗', '✓'],
]:
    row = t.add_row().cells
    for i, v in enumerate(row_data):
        row[i].text = v

add_heading('2.4 Neurosymbolic AI', level=2)
doc.add_paragraph(
    "Symbolic AI represents knowledge as explicit structures — interpretable but brittle. "
    "Neural AI learns from data — generalizes well but is opaque. In ExplainHire: the ontology "
    "graph handles structured skill reasoning (symbolic); SBERT handles semantic understanding "
    "(neural); XGBoost learns optimal signal fusion; SHAP provides post-hoc explanation. "
    "Every decision has a statistically learned justification (SHAP values) and a structural "
    "justification (ontology traversal paths + evidence sentences)."
)

add_heading('2.5 SHAP Explainability', level=2)
doc.add_paragraph(
    "SHAP attributes each feature's contribution to a prediction via Shapley values from "
    "cooperative game theory: φᵢ = average marginal contribution of feature i across all "
    "possible feature orderings. SHAP TreeExplainer computes exact Shapley values for XGBoost "
    "in polynomial time."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Chapter 3: SYSTEM ARCHITECTURE', level=1)

add_heading('3.1 Neurosymbolic Architecture Overview', level=2)
add_placeholder('fig3_architecture — Neurosymbolic Architecture Diagram')
doc.add_paragraph(
    "Three independent signals (symbolic graph, neural SBERT, structural) are computed in "
    "parallel and fused by XGBoost. See Figure 3.1."
)

add_heading('3.2 Design Principles', level=2)
add_numbered('Separation of concerns: each pipeline layer has a single, well-defined responsibility.')
add_numbered('Signal independence: the three signals are computed independently before fusion.')
add_numbered('Explainability by design: every component contributing to a decision also contributes to its explanation.')
add_numbered('Graceful degradation: meaningful output even when individual signals return edge-case values.')

add_heading('3.3 Pipeline Overview', level=2)
add_placeholder('fig1_pipeline — Seven-Layer Pipeline Diagram')

add_heading('3.4 L1 — Input Validation', level=2)
doc.add_paragraph(
    "Validates file type by magic bytes (PDF/DOCX only), enforces 5 MB size limit, "
    "detects and rejects scanned image PDFs, enforces minimum content length, and validates "
    "JD non-emptiness. Any failure returns an HTTP error with user-facing explanation."
)

add_heading('3.5 L2 — Parsing', level=2)
doc.add_paragraph(
    "PDF text extracted via PyMuPDF; DOCX via python-docx. From both resume and JD, the "
    "parser extracts: (a) Skills — vocabulary scan over ontology nodes + alias table; "
    "(b) YOE — regex patterns matching '3 years', '3+ years', '2 to 4 years'; "
    "(c) Education Level — degree keywords mapped to ranks 1–5; "
    "(d) Section Coverage — presence of Skills, Experience, Projects, Education headers."
)

add_heading('3.6 L3 — Normalization', level=2)
doc.add_paragraph(
    "The alias table (800+ entries) maps surface forms to canonical ontology node identifiers: "
    "'nodejs' → 'node.js', 'ML' → 'machine_learning', 'PyTorch' → 'pytorch'. "
    "Built by manual review of all 144 real resumes."
)

add_heading('3.7 L4 — Matching Layer', level=2)

add_heading('3.7.1 Graph Matcher (Symbolic AI)', level=3)
add_placeholder('fig4_ontology — Skill Ontology Subgraph')
doc.add_paragraph(
    "The skill ontology is a 434-node directed graph G=(V,E) of IS-A relationships, "
    "built from O*NET and manually extended for Indian tech hiring contexts. "
    "For each JD skill, the best-matching resume skill is found using six-level weights:"
)
add_bullet('1.0 — exact match')
add_bullet('0.9 — alias match')
add_bullet('0.7 — child to parent (resume skill is more specific)')
add_bullet('0.5 — parent to child (resume skill is more general)')
add_bullet('0.4 — sibling (common ancestor)')
add_bullet('0.3 — two-hop traversal')
doc.add_paragraph(
    "GraphScore(R,J) = mean of best match weights across all JD skills. "
    "5 features: {graph_score, coverage, exact_matches, partial_matches, no_matches}."
)

add_heading('3.7.2 Semantic Matcher (Neural AI)', level=3)
doc.add_paragraph(
    "Uses all-MiniLM-L6-v2 SBERT (6 transformer layers, 384-dim embeddings). "
    "SBERT_Score = cosine similarity between resume and JD embeddings. "
    "Long texts truncated to last 2,000 chars, prioritizing skills/projects sections. "
    "Fine-tuned with CosineSimilarityLoss on 887 real pairs for 4 epochs (loss=0.069). "
    "Contributes 1 feature: {sbert_score}."
)

add_heading('3.7.3 Structural Matcher', level=3)
doc.add_paragraph("YOE_Score = max(0, 1 - max(0, YOE_req - YOE_resume) / (YOE_req + 1))")
doc.add_paragraph("Edu_Score = 1.0 if rank_resume ≥ rank_req; 0.5 if one rank below; 0.0 otherwise.")
doc.add_paragraph("Section_Score = detected sections / 4.")
doc.add_paragraph("Contributes 8 features: {structural_score, yoe_score, edu_score, section_score, resume_yoe, required_yoe, resume_edu_rank, required_edu_rank}.")

add_heading('3.8 L5 — Classification and Explainability', level=2)
doc.add_paragraph(
    "14-feature vector x = [g1..g5 (graph), s1 (SBERT), t1..t8 (structural)] is fed to XGBoost. "
    "XGBoost builds 200 decision trees, each correcting the previous: ŷ(t) = ŷ(t-1) + η·f_t(x). "
    "Selected for: best tabular performance at small dataset sizes, native SHAP compatibility, "
    "and CPU-only training."
)
add_placeholder('fig5_shap — SHAP Feature Attribution Bar Chart')
doc.add_paragraph(
    "SHAP TreeExplainer computes Shapley values φᵢ for each of 14 features: f(x) = φ₀ + Σφᵢ. "
    "Positive φᵢ pushes toward Match; negative toward No Match."
)

add_heading('3.9 L6 — Explainer', level=2)
doc.add_paragraph(
    "Evidence Mapper: for each matched JD skill, finds the most relevant resume sentence "
    "via TF-IDF weighted token overlap. User sees: matched skill, ontological distance, "
    "and the specific resume sentence as evidence."
)
doc.add_paragraph(
    "Skill Gap Ranker: for each unmatched skill, computes Δⱼ = P_XGB(x'_j) - P_XGB(x), "
    "where x'_j hypothetically adds the missing skill. Gaps ranked by Δⱼ with learning resource links."
)

add_heading('3.10 L7 — Web Application', level=2)
doc.add_paragraph(
    "Flask application with: resume upload (PDF/DOCX drag-and-drop), JD text input, "
    "results dashboard (verdict, score cards, skill breakdown with evidence, skill gap list, "
    "SHAP bar chart), match history (SQLite/Flask-SQLAlchemy), user authentication (Flask-Login, "
    "bcrypt), interactive 434-node ontology visualizer (vis-network), dark/light mode toggle."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — DATASET AND TRAINING
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Chapter 4: DATASET AND TRAINING', level=1)

add_heading('4.1 Data Sources', level=2)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
hdr = t.rows[0].cells
for i, h in enumerate(['Source', 'Rows', 'Label Type', 'Notes']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ['Synthetic pairs', '167', 'Manual', 'Edge cases and boundary conditions'],
    ['Kaggle Resume Dataset', '800', 'Category-based', '24 job categories'],
    ['Friends dataset', '887', 'Score-based auto-label', '144 resumes × 17 real JDs'],
    ['Total', '1,854', '', ''],
]:
    row = t.add_row().cells
    for i, v in enumerate(row_data):
        row[i].text = v

doc.add_paragraph()
doc.add_paragraph(
    "Friends dataset: 144 real B.Tech resumes collected from our peer network, matched against "
    "17 real industry JDs from Naukri/LinkedIn covering Full Stack, Data Science, Backend, "
    "Flutter, ML Engineering, Golang, DevOps, and Android roles."
)

add_heading('4.2 Labeling Methodology', level=2)
doc.add_paragraph(
    "Auto-labeling rule for 2,448 friends pairs — requires both signals to agree strongly:\n"
    "  label = 1 if GraphScore > 0.4 AND SBERT_Score > 0.4\n"
    "  label = 0 if GraphScore < 0.3 AND SBERT_Score < 0.3\n"
    "  label = -1 (excluded) otherwise\n"
    "Result: 636 positive, 251 negative (887 total). 1,561 mixed pairs excluded to avoid label noise."
)

add_heading('4.3 SBERT Fine-Tuning', level=2)
doc.add_paragraph(
    "Base model: all-MiniLM-L6-v2. Loss: CosineSimilarityLoss. Epochs: 4. Batch size: 16. "
    "Warmup: 88 steps. Optimizer: AdamW. Final training loss: 0.069."
)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(['Configuration', 'CV F1', 'Full-data F1']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ['Base SBERT + 967 rows', '0.899', '0.899'],
    ['Fine-tuned SBERT + 1,854 rows', '0.9045', '0.951'],
]:
    row = t.add_row().cells
    for i, v in enumerate(row_data): row[i].text = v

add_heading('4.4 XGBoost Training', level=2)
doc.add_paragraph("14-feature vectors for all 1,854 pairs extracted and cached in data/annotation.csv.")
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
for i, h in enumerate(['Parameter', 'Value']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ['n_estimators', '200'], ['max_depth', '6'],
    ['learning_rate', '0.1'], ['tree_method', 'hist (CPU-optimized)'],
    ['random_state', '42'],
]:
    row = t.add_row().cells
    for i, v in enumerate(row_data): row[i].text = v
doc.add_paragraph()
doc.add_paragraph(
    "5-fold stratified cross-validation preserves the 71.7%/28.3% positive/negative class ratio. "
    "CV F1 = 0.9045 is the primary honest performance estimate."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 — EXPERIMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Chapter 5: EXPERIMENTS AND RESULTS', level=1)

add_heading('5.1 Ablation Study', level=2)
add_placeholder('fig2_ablation — Ablation Study Bar Chart')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
for i, h in enumerate(['Configuration', 'Macro F1', 'Δ vs. prev.', 'Interpretation']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ['TF-IDF baseline', '0.037', '—', 'Keyword matching fails'],
    ['SBERT-only', '0.484', '+44.7%', 'Semantics help; global sim imprecise'],
    ['Graph-only', '0.747', '+26.3%', 'Strongest single signal'],
    ['Graph + SBERT', '0.832', '+8.5%', 'Symbolic + neural complementary'],
    ['Full pipeline', '0.951', '+11.9%', 'Structural adds independent value'],
]:
    row = t.add_row().cells
    for i, v in enumerate(row_data): row[i].text = v

doc.add_paragraph()
doc.add_paragraph(
    "The graph-only result (F1=0.747) is the most informative: structured skill matching "
    "over an ontology substantially outperforms both keyword matching and global semantic "
    "similarity, validating the core architectural decision. The full pipeline (F1=0.951) "
    "confirms all three signals contribute independently."
)

add_heading('5.2 ResumeAtlas Benchmark', level=2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(['Method', 'Macro F1', 'vs. Heakl baseline']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ['TF-IDF + LR (Heakl et al.)', '0.610', '—'],
    ['SBERT + LR', '0.706', '+9.6%'],
    ['SBERT + XGBoost (ExplainHire)', '0.720', '+11.06%'],
]:
    row = t.add_row().cells
    for i, v in enumerate(row_data): row[i].text = v

add_heading('5.3 Comparison with Prior Work', level=2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
for i, h in enumerate(['System', 'Method', 'F1', 'Dataset', 'Explainable']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ['PJFNN', 'CNN joint repr.', '0.800', 'Baidu (private)', 'No'],
    ['Li et al.', 'BERT + attention', '0.792', 'CRC (private)', 'No'],
    ['Bian et al.', 'Graph + BERT', '~0.78', 'Private', 'No'],
    ['conSultantBERT', 'SBERT Siamese', '0.749', 'Randstad (private)', 'No'],
    ['Heakl et al.', 'TF-IDF + XGBoost', '0.610', 'ResumeAtlas', 'No'],
    ['ExplainHire (full)', 'Neurosymbolic', '0.951', 'This work', 'Yes'],
    ['ExplainHire (atlas)', 'SBERT + XGBoost', '0.720', 'ResumeAtlas', 'Yes'],
]:
    row = t.add_row().cells
    for i, v in enumerate(row_data): row[i].text = v
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 — DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Chapter 6: DISCUSSION', level=1)

add_heading('6.1 Signal Importance', level=2)
doc.add_paragraph(
    "Signal hierarchy: Graph (0.747) > SBERT (0.484) > Structural (+11.9% at margin). "
    "The graph is most valuable because resume–job matching is fundamentally a structured "
    "coverage problem — a knowledge graph is the natural representation for hierarchical skill "
    "coverage checking. The 6-level traversal credits candidates with related skills exactly "
    "as a human reviewer would. SBERT captures contextual alignment the graph cannot, but "
    "global similarity is imprecise for structured skill coverage. Structural features add "
    "orthogonal information: experience and education cannot be detected by skill signals alone."
)

add_heading('6.2 LLM Comparison', level=2)
doc.add_paragraph("ExplainHire advantages over LLM-based resume screening:")
add_bullet("Consistency: identical input always produces identical output — essential for fair employment screening.")
add_bullet("Cost and speed: milliseconds at zero marginal cost vs. expensive per-token LLM inference at scale.")
add_bullet("Structured explainability: SHAP values and ontology traversal paths constitute an audit trail.")
add_bullet("Regulatory compliance: feature-level attribution is auditable under EU AI Act and similar frameworks.")

add_heading('6.3 Limitations', level=2)
add_bullet("Dataset scope: biased toward Indian B.Tech students in software engineering; non-technical domains unevaluated.")
add_bullet("Auto-labeling noise: score-based thresholding may introduce label errors despite conservative thresholds.")
add_bullet("Ontology gaps: JD skills absent from the 434-node graph always score 0 regardless of resume content.")
add_bullet("YOE extraction: regex may underestimate implicit experience statements.")

add_heading('6.4 Future Work', level=2)
add_bullet("Ontology expansion to non-technical domains (finance, law, healthcare).")
add_bullet("Active learning using counterfactual gap ranking to prioritize manual labeling.")
add_bullet("Multi-lingual resume support (Hindi, regional Indian languages).")
add_bullet("Multi-candidate ranking extending binary classification to scored applicant pools.")
add_bullet("LLM integration as a fourth signal for narrative understanding.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Chapter 7: CONCLUSION', level=1)
doc.add_paragraph(
    "This report presented ExplainHire, a neurosymbolic AI system for explainable resume–job "
    "description matching. The seven-layer pipeline combines a 434-node symbolic skill ontology "
    "graph, a domain-adapted SBERT model fine-tuned on 887 real pairs, and structural features "
    "fused by XGBoost with SHAP explainability and sentence-level evidence mapping."
)
doc.add_paragraph("Key findings:")
add_numbered("All three signals are necessary. TF-IDF: F1=0.037; full pipeline: F1=0.951.")
add_numbered("The symbolic ontology graph is the most valuable single component (F1=0.747), outperforming neural SBERT alone (F1=0.484).")
add_numbered("Structural features add +11.9% F1 at the margin, demonstrating independent predictive power.")
add_numbered("ExplainHire generalizes to public benchmarks: F1=0.720 on ResumeAtlas (+11.06% over reported baseline).")
add_numbered("Explainability is achieved without sacrificing accuracy — ExplainHire is simultaneously the most accurate and the only explainable system among all comparable published works.")
doc.add_paragraph(
    "ExplainHire demonstrates that the neurosymbolic paradigm is a productive approach for "
    "high-stakes NLP tasks requiring both accuracy and accountability. Similar architectures "
    "could extend to medical diagnosis, legal document review, and other domains where decisions "
    "must be both accurate and auditable."
)
doc.add_page_break()

# ── Appendix: File Structure ───────────────────────────────────────────────────
add_heading('Appendix A: Project File Structure', level=1)
doc.add_paragraph("""
ExplainHire/
|-- app/                         Flask web application
|   |-- __init__.py
|   |-- models.py                SQLAlchemy DB models
|   `-- routes/
|       |-- auth.py              /login /register /logout
|       |-- match.py             /match (main pipeline)
|       |-- history.py           /history /result/<id>
|       `-- ontology_view.py     /ontology
|-- pipeline/
|   |-- l1_input/
|   |-- l2_parse/
|   |-- l3_normalize/
|   |-- l4_match/
|   |   |-- graph_matcher.py
|   |   |-- semantic_matcher.py
|   |   |-- structural_matcher.py
|   |   `-- finetune_sbert.py
|   |-- l5_classify/
|   |   |-- trainer.py
|   |   `-- predictor.py
|   |-- l6_explain/
|   |   |-- evidence_mapper.py
|   |   `-- skill_gap.py
|   `-- l7_present/
|       `-- orchestrator.py
|-- ontology/
|   |-- skill_ontology.gpickle   434-node NetworkX graph
|   |-- alias_table.json         800+ surface forms
|   `-- learning_resources.json
|-- data/
|   |-- raw/friends/             144 real B.Tech resumes
|   |-- raw/jds/                 17 real industry JDs
|   |-- annotation.csv           1,854 labeled pairs
|   `-- friends_rows.csv
|-- evaluation/
|   |-- ablation.py
|   |-- baselines.py
|   `-- benchmark_resumeatlas.py
|-- models/
|   |-- xgb_matcher.pkl
|   `-- sbert_finetuned/
|-- config.py
|-- requirements.txt
`-- run.py
""", style='No Spacing')
doc.add_paragraph()

# ── Appendix: Results Summary ──────────────────────────────────────────────────
add_heading('Appendix B: Key Numerical Results', level=1)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
for i, h in enumerate(['Metric', 'Value']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in [
    ('Pipeline layers', '7'), ('Ontology nodes', '434'),
    ('Alias table entries', '800+'), ('Total training pairs', '1,854'),
    ('SBERT fine-tuning pairs', '887'), ('SBERT fine-tuning loss', '0.069'),
    ('XGBoost trees', '200'), ('Feature vector size', '14'),
    ('Cross-validated F1', '0.9045'), ('Full-dataset F1', '0.951'),
    ('Best single-signal F1 (Graph-only)', '0.747'),
    ('Improvement over best baseline', '+11.86%'),
    ('TF-IDF baseline F1', '0.037'),
    ('ResumeAtlas F1', '0.720'),
    ('Improvement over Heakl et al.', '+11.06%'),
]:
    row = t.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
doc.add_page_break()

# ── Bibliography ───────────────────────────────────────────────────────────────
add_heading('REFERENCES', level=1)
refs = [
    "[1] R. Zhu et al., 'Person-Job Fit: Adapting the Right Talent for the Right Job with Joint Representation Learning,' ACM TMIS, vol. 9, no. 3, 2018. DOI: 10.1145/3234465",
    "[2] X. Li et al., 'BERT-based Neural Collaborative Filtering for Job-Resume Matching,' EMNLP 2020. DOI: 10.18653/v1/2020.emnlp-main.681",
    "[3] S. Bian et al., 'Learning to Match Jobs with Resumes from Sparse Interaction Data using Multi-View Co-Teaching Network,' CIKM 2020. DOI: 10.1145/3340531.3411929",
    "[4] D. Lavi et al., 'consultantBERT: Fine-tuned Siamese Sentence-BERT for Job Title Benchmarking,' arXiv:2109.06912, 2021.",
    "[5] A. Heakl et al., 'ResumeAtlas: Revisiting Resume Classification with Large-Scale Datasets and LLMs,' arXiv:2406.18125, 2024.",
    "[6] N. Reimers and I. Gurevych, 'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,' EMNLP 2019. DOI: 10.18653/v1/D19-1410",
    "[7] T. Chen and C. Guestrin, 'XGBoost: A Scalable Tree Boosting System,' KDD 2016. DOI: 10.1145/2939672.2939785",
    "[8] S. Lundberg and S.-I. Lee, 'A Unified Approach to Interpreting Model Predictions,' NeurIPS 2017. DOI: 10.48550/arXiv.1705.07874",
    "[9] National Center for O*NET Development, 'O*NET OnLine,' U.S. Dept. of Labor. https://www.onetonline.org",
]
for ref in refs:
    doc.add_paragraph(ref)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r'C:\Users\lalit\Desktop\ME\ExplainHire\paper\dtu_report.docx'
doc.save(out)
print(f'Saved: {out}')
